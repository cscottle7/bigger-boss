#!/usr/bin/env python3
"""
Bigger Boss Agent System - Professional Markdown to DOCX Converter
Converts markdown files to professionally formatted DOCX documents with Australian styling.
"""

import argparse
import logging
import os
import re
import sys
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import markdown
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class ProfessionalDocxConverter:
    """Professional DOCX converter with Australian business styling."""

    def __init__(self, style_profile: str = "professional"):
        """
        Initialise converter with specified style profile.

        Args:
            style_profile: Style profile to use ('professional', 'academic', 'marketing')
        """
        self.style_profile = style_profile
        self.style_config = self._get_style_config(style_profile)
        self.british_spellings = {
            # Common American to British spelling corrections
            r'\borganization\b': 'organisation',
            r'\borganizations\b': 'organisations',
            r'\borganizational\b': 'organisational',
            r'\brealize\b': 'realise',
            r'\brealizes\b': 'realises',
            r'\brealized\b': 'realised',
            r'\brealizing\b': 'realising',
            r'\boptimize\b': 'optimise',
            r'\boptimizes\b': 'optimises',
            r'\boptimized\b': 'optimised',
            r'\boptimizing\b': 'optimising',
            r'\banalyze\b': 'analyse',
            r'\banalyzes\b': 'analyses',
            r'\banalyzed\b': 'analysed',
            r'\banalyzing\b': 'analysing',
            r'\bcolor\b': 'colour',
            r'\bcolors\b': 'colours',
            r'\bcenter\b': 'centre',
            r'\bcentered\b': 'centred',
            r'\bcentering\b': 'centring',
            r'\bfavor\b': 'favour',
            r'\bfavors\b': 'favours',
            r'\bfavored\b': 'favoured',
            r'\bfavoring\b': 'favouring',
            r'\bhonor\b': 'honour',
            r'\bhonors\b': 'honours',
            r'\bhonored\b': 'honoured',
            r'\bhonoring\b': 'honouring',
            r'\bbehavior\b': 'behaviour',
            r'\bbehaviors\b': 'behaviours',
            r'\bbehavioral\b': 'behavioural',
            r'\btraveled\b': 'travelled',
            r'\btraveling\b': 'travelling',
            r'\bcanceled\b': 'cancelled',
            r'\bcanceling\b': 'cancelling',
            r'\bdefense\b': 'defence',
            r'\bdefenses\b': 'defences',
            r'\blicense\b': 'licence',  # Note: 'license' is verb form in British English
            r'\bpractice\b': 'practise',  # Note: 'practice' is noun form in British English
        }

    def _get_style_config(self, profile: str) -> Dict:
        """Get style configuration for specified profile."""
        configs = {
            "professional": {
                "font_name": "Calibri",
                "font_size": 11,
                "heading_font": "Calibri",
                "title_size": 16,
                "h1_size": 14,
                "h2_size": 13,
                "h3_size": 12,
                "h4_size": 11,
                "line_spacing": 1.15,
                "paragraph_spacing_before": 6,
                "paragraph_spacing_after": 6,
                "heading_spacing_before": 12,
                "heading_spacing_after": 6,
                "margins": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
                "colors": {
                    "heading": RGBColor(44, 62, 80),  # Professional dark blue
                    "text": RGBColor(0, 0, 0),        # Black
                    "accent": RGBColor(52, 152, 219),  # Blue accent
                }
            },
            "marketing": {
                "font_name": "Segoe UI",
                "font_size": 11,
                "heading_font": "Segoe UI Semibold",
                "title_size": 18,
                "h1_size": 15,
                "h2_size": 13,
                "h3_size": 12,
                "h4_size": 11,
                "line_spacing": 1.2,
                "paragraph_spacing_before": 8,
                "paragraph_spacing_after": 8,
                "heading_spacing_before": 14,
                "heading_spacing_after": 8,
                "margins": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
                "colors": {
                    "heading": RGBColor(231, 76, 60),   # Marketing red
                    "text": RGBColor(0, 0, 0),         # Black
                    "accent": RGBColor(46, 204, 113),   # Green accent
                }
            },
            "academic": {
                "font_name": "Times New Roman",
                "font_size": 12,
                "heading_font": "Times New Roman",
                "title_size": 14,
                "h1_size": 14,
                "h2_size": 13,
                "h3_size": 12,
                "h4_size": 12,
                "line_spacing": 2.0,  # Double spacing for academic
                "paragraph_spacing_before": 0,
                "paragraph_spacing_after": 0,
                "heading_spacing_before": 12,
                "heading_spacing_after": 6,
                "margins": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
                "colors": {
                    "heading": RGBColor(0, 0, 0),      # Black
                    "text": RGBColor(0, 0, 0),        # Black
                    "accent": RGBColor(0, 0, 0),       # Black
                }
            }
        }
        return configs.get(profile, configs["professional"])

    def _apply_british_spelling(self, text: str) -> str:
        """Convert American spellings to British spellings."""
        for american, british in self.british_spellings.items():
            text = re.sub(american, british, text, flags=re.IGNORECASE)
        return text

    def _setup_document_styles(self, doc: Document) -> None:
        """Set up custom styles for the document."""
        styles = doc.styles

        # Configure Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.style_config["font_name"]
        normal_font.size = Pt(self.style_config["font_size"])
        normal_font.color.rgb = self.style_config["colors"]["text"]

        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.line_spacing = self.style_config["line_spacing"]
        normal_paragraph.space_before = Pt(self.style_config["paragraph_spacing_before"])
        normal_paragraph.space_after = Pt(self.style_config["paragraph_spacing_after"])

        # Configure heading styles
        heading_configs = [
            ('Heading 1', 'h1_size', True),
            ('Heading 2', 'h2_size', True),
            ('Heading 3', 'h3_size', True),
            ('Heading 4', 'h4_size', False),
        ]

        for style_name, size_key, bold in heading_configs:
            if style_name in styles:
                heading_style = styles[style_name]
            else:
                heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

            heading_font = heading_style.font
            heading_font.name = self.style_config["heading_font"]
            heading_font.size = Pt(self.style_config[size_key])
            heading_font.color.rgb = self.style_config["colors"]["heading"]
            heading_font.bold = bold

            heading_paragraph = heading_style.paragraph_format
            heading_paragraph.space_before = Pt(self.style_config["heading_spacing_before"])
            heading_paragraph.space_after = Pt(self.style_config["heading_spacing_after"])

        # Create custom styles
        self._create_custom_styles(styles)

    def _create_custom_styles(self, styles) -> None:
        """Create custom styles for special elements."""

        # Citation style
        citation_style = styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)
        citation_font = citation_style.font
        citation_font.name = self.style_config["font_name"]
        citation_font.size = Pt(self.style_config["font_size"] - 1)
        citation_font.italic = True
        citation_font.color.rgb = RGBColor(128, 128, 128)

        citation_paragraph = citation_style.paragraph_format
        citation_paragraph.left_indent = Inches(0.5)
        citation_paragraph.space_before = Pt(3)
        citation_paragraph.space_after = Pt(3)

        # Quote style
        quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
        quote_font = quote_style.font
        quote_font.name = self.style_config["font_name"]
        quote_font.size = Pt(self.style_config["font_size"])
        quote_font.italic = True

        quote_paragraph = quote_style.paragraph_format
        quote_paragraph.left_indent = Inches(0.5)
        quote_paragraph.right_indent = Inches(0.5)
        quote_paragraph.space_before = Pt(6)
        quote_paragraph.space_after = Pt(6)

        # Code style
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = "Consolas"
        code_font.size = Pt(self.style_config["font_size"] - 1)
        code_font.color.rgb = RGBColor(102, 102, 102)

        code_paragraph = code_style.paragraph_format
        code_paragraph.left_indent = Inches(0.5)
        code_paragraph.space_before = Pt(6)
        code_paragraph.space_after = Pt(6)

    def _set_document_margins(self, doc: Document) -> None:
        """Set document margins."""
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(self.style_config["margins"]["top"])
            section.bottom_margin = Cm(self.style_config["margins"]["bottom"])
            section.left_margin = Cm(self.style_config["margins"]["left"])
            section.right_margin = Cm(self.style_config["margins"]["right"])

    def _parse_markdown(self, markdown_content: str) -> BeautifulSoup:
        """Parse markdown content to HTML and then to BeautifulSoup."""
        # Configure markdown extensions for better parsing
        md = markdown.Markdown(
            extensions=[
                'extra',
                'codehilite',
                'toc',
                'tables',
                'fenced_code',
                'nl2br',
            ],
            extension_configs={
                'codehilite': {
                    'css_class': 'highlight'
                },
                'toc': {
                    'anchorlink': True
                }
            }
        )

        # Apply British spelling corrections
        markdown_content = self._apply_british_spelling(markdown_content)

        # Convert markdown to HTML
        html_content = md.convert(markdown_content)

        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        return soup

    def _process_element(self, element, doc: Document) -> None:
        """Process a BeautifulSoup element and add it to the document."""

        if element.name == 'h1':
            self._add_heading(doc, element.get_text().strip(), 1)
        elif element.name == 'h2':
            self._add_heading(doc, element.get_text().strip(), 2)
        elif element.name == 'h3':
            self._add_heading(doc, element.get_text().strip(), 3)
        elif element.name == 'h4':
            self._add_heading(doc, element.get_text().strip(), 4)
        elif element.name == 'p':
            self._add_paragraph(doc, element)
        elif element.name == 'ul':
            self._add_list(doc, element, bullet=True)
        elif element.name == 'ol':
            self._add_list(doc, element, bullet=False)
        elif element.name == 'blockquote':
            self._add_quote(doc, element)
        elif element.name == 'pre':
            self._add_code_block(doc, element)
        elif element.name == 'table':
            self._add_table(doc, element)
        elif element.name in ['hr']:
            self._add_horizontal_rule(doc)
        elif isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                doc.add_paragraph(text)

    def _add_heading(self, doc: Document, text: str, level: int) -> None:
        """Add a heading to the document."""
        heading_styles = {
            1: 'Heading 1',
            2: 'Heading 2',
            3: 'Heading 3',
            4: 'Heading 4',
        }

        style_name = heading_styles.get(level, 'Heading 1')
        heading = doc.add_heading(text, level=0)
        heading.style = style_name

    def _add_paragraph(self, doc: Document, element) -> None:
        """Add a paragraph with inline formatting."""
        paragraph = doc.add_paragraph()

        # Check if this is a citation paragraph
        text_content = element.get_text()
        if text_content.startswith('**Source:**') or text_content.startswith('Source:'):
            paragraph.style = 'Citation'

        # Process inline elements
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip():
                    run = paragraph.add_run(text)
            elif child.name == 'strong' or child.name == 'b':
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name == 'em' or child.name == 'i':
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == 'code':
                run = paragraph.add_run(child.get_text())
                run.font.name = 'Consolas'
                run.font.size = Pt(self.style_config["font_size"] - 1)
            elif child.name == 'a':
                run = paragraph.add_run(child.get_text())
                run.font.color.rgb = self.style_config["colors"]["accent"]
                run.underline = True
            else:
                run = paragraph.add_run(child.get_text())

        # Remove empty paragraphs
        if not paragraph.text.strip():
            doc._body._remove_element(paragraph._element)

    def _add_list(self, doc: Document, element, bullet: bool = True) -> None:
        """Add a bulleted or numbered list."""
        for li in element.find_all('li', recursive=False):
            paragraph = doc.add_paragraph(style='List Bullet' if bullet else 'List Number')

            # Process inline elements in list items
            for child in li.children:
                if isinstance(child, NavigableString):
                    text = str(child)
                    if text.strip():
                        run = paragraph.add_run(text)
                elif child.name == 'strong' or child.name == 'b':
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'em' or child.name == 'i':
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                else:
                    run = paragraph.add_run(child.get_text())

    def _add_quote(self, doc: Document, element) -> None:
        """Add a blockquote."""
        paragraph = doc.add_paragraph(element.get_text(), style='Quote')

    def _add_code_block(self, doc: Document, element) -> None:
        """Add a code block."""
        code_text = element.get_text()
        paragraph = doc.add_paragraph(code_text, style='Code')

    def _add_table(self, doc: Document, element) -> None:
        """Add a table to the document."""
        rows = element.find_all('tr')
        if not rows:
            return

        # Get the maximum number of columns
        max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)

        # Create table
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'

        # Populate table
        for row_idx, tr in enumerate(rows):
            cells = tr.find_all(['th', 'td'])
            for col_idx, cell in enumerate(cells):
                if col_idx < max_cols:
                    table_cell = table.cell(row_idx, col_idx)
                    table_cell.text = cell.get_text().strip()

                    # Apply header styling for th elements
                    if cell.name == 'th':
                        for paragraph in table_cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    def _add_horizontal_rule(self, doc: Document) -> None:
        """Add a horizontal rule (line break)."""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.LINE)

    def convert_file(self, input_path: str, output_path: Optional[str] = None) -> str:
        """
        Convert a markdown file to DOCX format.

        Args:
            input_path: Path to input markdown file
            output_path: Path to output DOCX file (optional)

        Returns:
            Path to the created DOCX file
        """
        input_file = Path(input_path)

        if not input_file.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")

        if not input_file.suffix.lower() == '.md':
            raise ValueError(f"Input file must be a markdown file (.md): {input_path}")

        # Determine output path
        if output_path is None:
            output_path = input_file.with_suffix('.docx')
        else:
            output_path = Path(output_path)

        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            # Read markdown content
            with open(input_file, 'r', encoding='utf-8') as f:
                markdown_content = f.read()

            logger.info(f"Converting {input_file} to {output_path}")

            # Create document
            doc = Document()

            # Set up styles and margins
            self._setup_document_styles(doc)
            self._set_document_margins(doc)

            # Parse markdown content
            soup = self._parse_markdown(markdown_content)

            # Process all elements
            for element in soup.children:
                if element.name:  # Skip text nodes at root level
                    self._process_element(element, doc)

            # Save document
            doc.save(str(output_path))

            logger.info(f"Successfully converted {input_file} to {output_path}")
            return str(output_path)

        except Exception as e:
            logger.error(f"Error converting {input_file}: {str(e)}")
            raise

    def convert_content(self, markdown_content: str, output_path: str) -> str:
        """
        Convert markdown content directly to DOCX format.

        Args:
            markdown_content: Markdown content as string
            output_path: Path to output DOCX file

        Returns:
            Path to the created DOCX file
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            logger.info(f"Converting markdown content to {output_path}")

            # Create document
            doc = Document()

            # Set up styles and margins
            self._setup_document_styles(doc)
            self._set_document_margins(doc)

            # Parse markdown content
            soup = self._parse_markdown(markdown_content)

            # Process all elements
            for element in soup.children:
                if element.name:  # Skip text nodes at root level
                    self._process_element(element, doc)

            # Save document
            doc.save(str(output_path))

            logger.info(f"Successfully converted content to {output_path}")
            return str(output_path)

        except Exception as e:
            logger.error(f"Error converting content: {str(e)}")
            raise


def main():
    """Main function for command-line usage."""
    parser = argparse.ArgumentParser(
        description='Convert Markdown files to professional DOCX documents'
    )
    parser.add_argument(
        'input_file',
        help='Path to input markdown file'
    )
    parser.add_argument(
        '-o', '--output',
        help='Path to output DOCX file (optional)'
    )
    parser.add_argument(
        '--style',
        choices=['professional', 'marketing', 'academic'],
        default='professional',
        help='Style profile to use (default: professional)'
    )
    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose logging'
    )

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    try:
        converter = ProfessionalDocxConverter(style_profile=args.style)
        output_file = converter.convert_file(args.input_file, args.output)
        print(f"✅ Successfully converted to: {output_file}")

    except Exception as e:
        print(f"❌ Error: {str(e)}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()