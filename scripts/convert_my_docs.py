#!/usr/bin/env python3
"""
Bigger Boss Agent System - Document Converter
Converts markdown files to .docx format for client delivery.

Features:
- Converts markdown to properly formatted Word documents
- Preserves formatting, headings, and structure
- Handles British English formatting
- Batch conversion support
- Error handling and logging

Usage:
    python scripts/convert_my_docs.py file.md
    python scripts/convert_my_docs.py folder_path --batch
    python scripts/convert_my_docs.py --test
"""

import argparse
import json
import logging
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import io

# Fix Windows Unicode encoding issues
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

try:
    import pandoc
    from pandoc.types import *
    PANDOC_AVAILABLE = True
except ImportError:
    try:
        import pypandoc
        PANDOC_AVAILABLE = True
        PANDOC_BACKEND = 'pypandoc'
    except ImportError:
        PANDOC_AVAILABLE = False
        logger.warning("Neither pandoc nor pypandoc available - falling back to basic conversion")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger.warning("python-docx not available - basic text conversion only")


class MarkdownToDocxConverter:
    """
    Converts markdown files to properly formatted Word documents.
    """

    def __init__(self):
        self.conversion_stats = {
            'total_files': 0,
            'successful': 0,
            'failed': 0,
            'skipped': 0
        }

    def setup_document_styles(self, doc: Document):
        """Set up custom styles for the document."""
        if not PYTHON_DOCX_AVAILABLE:
            return

        try:
            # Add custom heading styles if they don't exist
            styles = doc.styles

            # Main heading style
            try:
                heading_style = styles['Heading 1']
                heading_style.font.size = Pt(16)
                heading_style.font.bold = True
            except KeyError:
                # Create if doesn't exist
                heading_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
                heading_style.font.size = Pt(16)
                heading_style.font.bold = True

            # Subtitle style
            try:
                subtitle_style = styles['Heading 2']
                subtitle_style.font.size = Pt(14)
                subtitle_style.font.bold = True
            except KeyError:
                subtitle_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
                subtitle_style.font.size = Pt(14)
                subtitle_style.font.bold = True

            # Normal text style
            try:
                normal_style = styles['Normal']
                normal_style.font.size = Pt(11)
                normal_style.font.name = 'Calibri'
            except KeyError:
                pass

        except Exception as e:
            logger.warning(f"Could not set up document styles: {e}")

    def convert_with_pandoc(self, input_file: str, output_file: str) -> bool:
        """Convert using pandoc (preferred method)."""
        try:
            if 'pypandoc' in globals():
                # Use pypandoc
                output = pypandoc.convert_file(
                    input_file,
                    'docx',
                    outputfile=output_file,
                    extra_args=[
                        '--reference-doc=scripts/templates/reference.docx' if
                        Path('scripts/templates/reference.docx').exists() else None
                    ]
                )
                return True
            else:
                # Use pandoc directly
                import subprocess
                cmd = [
                    'pandoc',
                    input_file,
                    '-o', output_file,
                    '--from=markdown',
                    '--to=docx'
                ]

                # Add reference document if available
                ref_doc = Path('scripts/templates/reference.docx')
                if ref_doc.exists():
                    cmd.extend(['--reference-doc', str(ref_doc)])

                result = subprocess.run(cmd, capture_output=True, text=True)
                return result.returncode == 0

        except Exception as e:
            logger.error(f"Pandoc conversion failed: {e}")
            return False

    def parse_inline_markdown(self, text: str, paragraph):
        """Parse inline markdown formatting (bold, italic, code)."""
        import re

        # Pattern to match **bold**, *italic*, or `code`
        pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Italic text
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Code text
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                # Regular text
                paragraph.add_run(part)

    def convert_with_python_docx(self, input_file: str, output_file: str) -> bool:
        """Convert using python-docx with proper inline markdown formatting."""
        if not PYTHON_DOCX_AVAILABLE:
            return False

        try:
            # Read markdown file
            with open(input_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # Create new document
            doc = Document()
            self.setup_document_styles(doc)

            # Parse markdown content line by line
            lines = content.split('\n')
            in_list = False

            for line in lines:
                stripped = line.strip()

                if not stripped:
                    # Empty line - add paragraph break
                    doc.add_paragraph()
                    in_list = False
                    continue

                # Check for headings
                if stripped.startswith('# '):
                    # H1 heading
                    heading_text = stripped[2:]
                    heading = doc.add_heading(level=1)
                    self.parse_inline_markdown(heading_text, heading)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    in_list = False

                elif stripped.startswith('## '):
                    # H2 heading
                    heading_text = stripped[3:]
                    heading = doc.add_heading(level=2)
                    self.parse_inline_markdown(heading_text, heading)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    in_list = False

                elif stripped.startswith('### '):
                    # H3 heading
                    heading_text = stripped[4:]
                    heading = doc.add_heading(level=3)
                    self.parse_inline_markdown(heading_text, heading)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    in_list = False

                elif stripped.startswith('- ') or stripped.startswith('* '):
                    # Bullet point
                    bullet_text = stripped[2:]
                    p = doc.add_paragraph(style='List Bullet')
                    self.parse_inline_markdown(bullet_text, p)
                    in_list = True

                elif stripped.startswith('1. ') or (len(stripped) > 2 and stripped[0].isdigit() and stripped[1:].startswith('. ')):
                    # Numbered list
                    # Find where the actual text starts (after number and dot)
                    idx = stripped.find('. ')
                    if idx > 0:
                        list_text = stripped[idx+2:]
                        p = doc.add_paragraph(style='List Number')
                        self.parse_inline_markdown(list_text, p)
                        in_list = True

                elif stripped.startswith('---') or stripped.startswith('***'):
                    # Horizontal rule
                    doc.add_paragraph()
                    p = doc.add_paragraph('_' * 50)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                    in_list = False

                else:
                    # Regular paragraph with inline formatting
                    p = doc.add_paragraph()
                    self.parse_inline_markdown(stripped, p)
                    in_list = False

            # Save document
            doc.save(output_file)
            return True

        except Exception as e:
            logger.error(f"Python-docx conversion failed: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False

    def convert_basic_text(self, input_file: str, output_file: str) -> bool:
        """Basic text conversion (last resort)."""
        try:
            # Read markdown content
            with open(input_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # Simple text cleaning
            content = content.replace('# ', '')
            content = content.replace('## ', '')
            content = content.replace('### ', '')
            content = content.replace('**', '')
            content = content.replace('*', '')
            content = content.replace('---', '\n' + '='*50 + '\n')

            # Save as .txt file (user can manually convert)
            txt_output = output_file.replace('.docx', '.txt')
            with open(txt_output, 'w', encoding='utf-8') as f:
                f.write(f"Converted from: {input_file}\n")
                f.write(f"Conversion date: {datetime.now().strftime('%d %B %Y')}\n")
                f.write('='*50 + '\n\n')
                f.write(content)

            logger.warning(f"Basic text conversion saved to: {txt_output}")
            logger.warning("Please manually convert to .docx format")
            return True

        except Exception as e:
            logger.error(f"Basic text conversion failed: {e}")
            return False

    def convert_file(self, input_file: str, output_file: Optional[str] = None) -> bool:
        """
        Convert a single markdown file to .docx format.

        Args:
            input_file: Path to input .md file
            output_file: Path to output .docx file (optional)

        Returns:
            True if conversion successful, False otherwise
        """
        input_path = Path(input_file)

        if not input_path.exists():
            logger.error(f"Input file does not exist: {input_file}")
            return False

        if input_path.suffix.lower() != '.md':
            logger.warning(f"Input file is not a markdown file: {input_file}")
            return False

        # Determine output file path
        if not output_file:
            output_file = str(input_path.with_suffix('.docx'))

        output_path = Path(output_file)

        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        logger.info(f"Converting: {input_file} -> {output_file}")

        # Try conversion methods in order of preference
        success = False

        if PANDOC_AVAILABLE and not success:
            logger.debug("Attempting pandoc conversion...")
            success = self.convert_with_pandoc(str(input_path), str(output_path))

        if not success and PYTHON_DOCX_AVAILABLE:
            logger.debug("Attempting python-docx conversion...")
            success = self.convert_with_python_docx(str(input_path), str(output_path))

        if not success:
            logger.debug("Attempting basic text conversion...")
            success = self.convert_basic_text(str(input_path), str(output_path))

        # Update statistics
        self.conversion_stats['total_files'] += 1
        if success:
            self.conversion_stats['successful'] += 1
            logger.info(f"✅ Successfully converted: {input_file}")
        else:
            self.conversion_stats['failed'] += 1
            logger.error(f"❌ Failed to convert: {input_file}")

        return success

    def convert_folder(self, folder_path: str, recursive: bool = True) -> Dict[str, bool]:
        """
        Convert all markdown files in a folder.

        Args:
            folder_path: Path to folder containing .md files
            recursive: Whether to search subfolders

        Returns:
            Dictionary mapping file paths to conversion success status
        """
        folder = Path(folder_path)

        if not folder.exists() or not folder.is_dir():
            logger.error(f"Folder does not exist: {folder_path}")
            return {}

        # Find all markdown files
        if recursive:
            md_files = list(folder.rglob("*.md"))
        else:
            md_files = list(folder.glob("*.md"))

        if not md_files:
            logger.warning(f"No markdown files found in: {folder_path}")
            return {}

        logger.info(f"Found {len(md_files)} markdown files to convert")

        results = {}
        for md_file in md_files:
            try:
                success = self.convert_file(str(md_file))
                results[str(md_file)] = success
            except Exception as e:
                logger.error(f"Error converting {md_file}: {e}")
                results[str(md_file)] = False

        return results

    def print_statistics(self):
        """Print conversion statistics."""
        stats = self.conversion_stats
        print("\n" + "="*50)
        print("CONVERSION STATISTICS")
        print("="*50)
        print(f"Total files processed: {stats['total_files']}")
        print(f"Successful conversions: {stats['successful']}")
        print(f"Failed conversions: {stats['failed']}")
        print(f"Skipped files: {stats['skipped']}")

        if stats['total_files'] > 0:
            success_rate = (stats['successful'] / stats['total_files']) * 100
            print(f"Success rate: {success_rate:.1f}%")

        print("="*50)

    def test_conversion_capability(self) -> bool:
        """Test conversion capabilities and report what's available."""
        print("\n" + "="*50)
        print("CONVERSION CAPABILITY TEST")
        print("="*50)

        capabilities = []

        if PANDOC_AVAILABLE:
            capabilities.append("✅ Pandoc conversion (preferred)")
        else:
            capabilities.append("❌ Pandoc conversion (not available)")

        if PYTHON_DOCX_AVAILABLE:
            capabilities.append("✅ Python-docx conversion (fallback)")
        else:
            capabilities.append("❌ Python-docx conversion (not available)")

        capabilities.append("✅ Basic text conversion (always available)")

        for capability in capabilities:
            print(capability)

        print("\nRecommendation:")
        if PANDOC_AVAILABLE:
            print("✅ Full conversion capability available")
            print("Install: pip install pandoc pypandoc")
            return True
        elif PYTHON_DOCX_AVAILABLE:
            print("⚠️  Limited conversion capability")
            print("Install pandoc for better results: pip install pandoc pypandoc")
            return True
        else:
            print("❌ Minimal conversion capability")
            print("Install: pip install pandoc pypandoc python-docx")
            return False


def main():
    """Main function for command-line usage."""
    parser = argparse.ArgumentParser(
        description='Convert markdown files to .docx format for client delivery'
    )

    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument(
        'input_path',
        nargs='?',
        help='Path to markdown file or folder to convert'
    )
    group.add_argument(
        '--test',
        action='store_true',
        help='Test conversion capabilities'
    )

    parser.add_argument(
        '--output',
        help='Output file path (for single file conversion)'
    )
    parser.add_argument(
        '--batch',
        action='store_true',
        help='Convert all .md files in the specified folder'
    )
    parser.add_argument(
        '--recursive',
        action='store_true',
        default=True,
        help='Search subfolders when using --batch (default: True)'
    )
    parser.add_argument(
        '--verbose',
        action='store_true',
        help='Enable verbose logging'
    )

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    converter = MarkdownToDocxConverter()

    try:
        if args.test:
            # Test mode
            converter.test_conversion_capability()
            return

        if not args.input_path:
            parser.print_help()
            return

        input_path = Path(args.input_path)

        if not input_path.exists():
            print(f"❌ Error: Path does not exist: {args.input_path}")
            sys.exit(1)

        if input_path.is_file():
            # Single file conversion
            if args.batch:
                print("❌ Error: --batch flag cannot be used with single file")
                sys.exit(1)

            success = converter.convert_file(str(input_path), args.output)
            converter.print_statistics()
            sys.exit(0 if success else 1)

        elif input_path.is_dir():
            # Folder conversion
            if not args.batch:
                print("❌ Error: Use --batch flag to convert folder contents")
                sys.exit(1)

            results = converter.convert_folder(str(input_path), args.recursive)
            converter.print_statistics()

            # Exit with success if any files were converted
            successful_conversions = sum(1 for success in results.values() if success)
            sys.exit(0 if successful_conversions > 0 else 1)

        else:
            print(f"❌ Error: Invalid path type: {args.input_path}")
            sys.exit(1)

    except KeyboardInterrupt:
        print("\n❌ Conversion cancelled by user")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Conversion error: {str(e)}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()