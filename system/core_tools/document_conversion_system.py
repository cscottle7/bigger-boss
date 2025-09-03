#!/usr/bin/env python3
"""
Document Conversion System for Marketing Analysis Reports
Converts Markdown reports to .docx format with rich text formatting
Integrates with Google Drive agents for automated file management
"""

import os
import re
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARNING] python-docx not available. Install with: pip install python-docx")

try:
    import markdown
    from markdown.extensions import tables, toc
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    print("[WARNING] markdown not available. Install with: pip install markdown")

class DocumentConverter:
    """Advanced document conversion system with rich formatting support"""
    
    def __init__(self):
        self.check_dependencies()
        self.template_styles = {
            'title': {'font_size': 18, 'bold': True, 'color': '2F4F4F'},
            'heading1': {'font_size': 16, 'bold': True, 'color': '4682B4'},
            'heading2': {'font_size': 14, 'bold': True, 'color': '708090'},
            'heading3': {'font_size': 12, 'bold': True, 'color': '696969'},
            'body': {'font_size': 11, 'color': '000000'},
            'table_header': {'font_size': 10, 'bold': True, 'color': '000000'},
            'table_cell': {'font_size': 10, 'color': '000000'},
            'code': {'font_size': 9, 'font_name': 'Courier New', 'color': '8B0000'},
            'quote': {'font_size': 11, 'italic': True, 'color': '4682B4'}
        }
    
    def check_dependencies(self):
        """Check if required dependencies are available"""
        missing_deps = []
        if not DOCX_AVAILABLE:
            missing_deps.append("python-docx")
        if not MARKDOWN_AVAILABLE:
            missing_deps.append("markdown")
        
        if missing_deps:
            raise ImportError(f"Missing required dependencies: {', '.join(missing_deps)}")
    
    def parse_markdown_content(self, markdown_text):
        """Parse markdown content and extract structured elements"""
        lines = markdown_text.split('\n')
        structured_content = []
        current_table = []
        in_code_block = False
        code_content = []
        
        for line in lines:
            # Handle code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End of code block
                    structured_content.append({
                        'type': 'code_block',
                        'content': '\n'.join(code_content),
                        'language': code_content[0] if code_content and code_content[0] in ['python', 'javascript', 'bash'] else 'text'
                    })
                    code_content = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                continue
            
            if in_code_block:
                code_content.append(line)
                continue
            
            # Handle tables
            if '|' in line and line.strip():
                if not current_table:
                    current_table = []
                current_table.append(line.strip())
                continue
            elif current_table:
                # End of table
                structured_content.append({
                    'type': 'table',
                    'content': self.parse_markdown_table(current_table)
                })
                current_table = []
            
            # Handle headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('# ').strip()
                structured_content.append({
                    'type': 'heading',
                    'level': level,
                    'content': text
                })
            # Handle lists
            elif line.strip().startswith(('- ', '* ', '+ ')):
                text = line.strip()[2:].strip()
                structured_content.append({
                    'type': 'bullet_list',
                    'content': text
                })
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                structured_content.append({
                    'type': 'numbered_list',
                    'content': text
                })
            # Handle quotes
            elif line.strip().startswith('>'):
                text = line.strip()[1:].strip()
                structured_content.append({
                    'type': 'quote',
                    'content': text
                })
            # Handle regular paragraphs
            elif line.strip():
                structured_content.append({
                    'type': 'paragraph',
                    'content': line.strip()
                })
            else:
                # Empty line
                structured_content.append({
                    'type': 'empty_line',
                    'content': ''
                })
        
        # Handle any remaining table
        if current_table:
            structured_content.append({
                'type': 'table',
                'content': self.parse_markdown_table(current_table)
            })
        
        return structured_content
    
    def parse_markdown_table(self, table_lines):
        """Parse markdown table into structured format"""
        if len(table_lines) < 2:
            return {'headers': [], 'rows': []}
        
        # Parse header
        headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        
        # Skip separator line (table_lines[1])
        rows = []
        for line in table_lines[2:]:
            if '|' in line:
                row = [cell.strip() for cell in line.split('|')[1:-1]]
                rows.append(row)
        
        return {'headers': headers, 'rows': rows}
    
    def apply_text_formatting(self, run, text):
        """Apply text formatting based on markdown syntax"""
        # Bold text
        if '**' in text:
            parts = text.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Odd indices are bold
                    run.bold = True
                    run.text = part
                    run.bold = False
                else:
                    run.text = part
        # Italic text
        elif '*' in text:
            parts = text.split('*')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Odd indices are italic
                    run.italic = True
                    run.text = part
                    run.italic = False
                else:
                    run.text = part
        # Code text
        elif '`' in text:
            parts = text.split('`')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Odd indices are code
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.text = part
                else:
                    run.text = part
        else:
            run.text = text
    
    def create_docx_document(self, structured_content, title="Marketing Analysis Report"):
        """Create a .docx document from structured content"""
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = "Marketing Analysis System"
        doc.core_properties.created = datetime.now()
        
        # Add document title
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation timestamp
        timestamp_paragraph = doc.add_paragraph()
        timestamp_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp_run = timestamp_paragraph.add_run(f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        timestamp_run.italic = True
        
        doc.add_paragraph()  # Empty line
        
        # Process structured content
        for element in structured_content:
            if element['type'] == 'heading':
                level = min(element['level'], 3)  # Limit to 3 levels
                doc.add_heading(element['content'], level)
            
            elif element['type'] == 'paragraph':
                paragraph = doc.add_paragraph()
                self.apply_text_formatting(paragraph.add_run(), element['content'])
            
            elif element['type'] == 'bullet_list':
                paragraph = doc.add_paragraph(element['content'], style='List Bullet')
            
            elif element['type'] == 'numbered_list':
                paragraph = doc.add_paragraph(element['content'], style='List Number')
            
            elif element['type'] == 'quote':
                paragraph = doc.add_paragraph()
                quote_run = paragraph.add_run(element['content'])
                quote_run.italic = True
                paragraph.paragraph_format.left_indent = Inches(0.5)
            
            elif element['type'] == 'code_block':
                paragraph = doc.add_paragraph()
                code_run = paragraph.add_run(element['content'])
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
                paragraph.paragraph_format.left_indent = Inches(0.5)
            
            elif element['type'] == 'table':
                self.add_table_to_doc(doc, element['content'])
            
            elif element['type'] == 'empty_line':
                doc.add_paragraph()
        
        return doc
    
    def add_table_to_doc(self, doc, table_data):
        """Add a formatted table to the document"""
        if not table_data['headers'] or not table_data['rows']:
            return
        
        # Create table
        table = doc.add_table(rows=1, cols=len(table_data['headers']))
        table.style = 'Table Grid'
        
        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(table_data['headers']):
            cell = header_row.cells[i]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(header)
            run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add data rows
        for row_data in table_data['rows']:
            row = table.add_row()
            for i, cell_data in enumerate(row_data):
                if i < len(row.cells):
                    row.cells[i].text = cell_data
        
        # Add spacing after table
        doc.add_paragraph()
    
    def hex_to_rgb(self, hex_color):
        """Convert hex color to RGB tuple"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    def convert_markdown_to_docx(self, markdown_file_path, output_dir=None):
        """Convert markdown file to .docx format"""
        if not os.path.exists(markdown_file_path):
            raise FileNotFoundError(f"Markdown file not found: {markdown_file_path}")
        
        # Read markdown content
        with open(markdown_file_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Extract title from filename or first heading
        title = Path(markdown_file_path).stem.replace('_', ' ').title()
        first_heading = re.search(r'^#\s+(.+)', markdown_content, re.MULTILINE)
        if first_heading:
            title = first_heading.group(1)
        
        # Parse markdown content
        structured_content = self.parse_markdown_content(markdown_content)
        
        # Create .docx document
        doc = self.create_docx_document(structured_content, title)
        
        # Determine output path
        if output_dir is None:
            output_dir = os.path.dirname(markdown_file_path)
        
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        output_filename = f"{Path(markdown_file_path).stem}.docx"
        output_path = os.path.join(output_dir, output_filename)
        
        # Save document
        doc.save(output_path)
        
        return output_path
    
    def batch_convert_directory(self, source_dir, output_dir=None, pattern="*.md"):
        """Convert all markdown files in a directory"""
        source_path = Path(source_dir)
        if output_dir is None:
            output_dir = source_dir
        
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        converted_files = []
        
        for md_file in source_path.glob(pattern):
            try:
                docx_file = self.convert_markdown_to_docx(str(md_file), str(output_path))
                converted_files.append({
                    'source': str(md_file),
                    'output': docx_file,
                    'status': 'success'
                })
                print(f"[SUCCESS] Converted: {md_file.name} -> {Path(docx_file).name}")
            except Exception as e:
                converted_files.append({
                    'source': str(md_file),
                    'output': None,
                    'status': 'error',
                    'error': str(e)
                })
                print(f"[ERROR] Failed to convert {md_file.name}: {e}")
        
        return converted_files

class GoogleDriveIntegration:
    """Integration with Google Drive agents for automated file management"""
    
    def __init__(self, converter):
        self.converter = converter
    
    def convert_and_upload(self, markdown_file, drive_folder="Marketing Reports"):
        """Convert markdown to .docx and upload to Google Drive"""
        try:
            # Convert to .docx
            docx_file = self.converter.convert_markdown_to_docx(markdown_file)
            
            # Upload to Google Drive (would integrate with google_drive_publisher agent)
            upload_result = self.upload_to_drive(docx_file, drive_folder)
            
            return {
                'status': 'success',
                'local_file': docx_file,
                'drive_file': upload_result.get('file_url', 'N/A'),
                'drive_id': upload_result.get('file_id', 'N/A')
            }
        
        except Exception as e:
            return {
                'status': 'error',
                'error': str(e)
            }
    
    def upload_to_drive(self, local_file, folder_name):
        """Upload file to Google Drive via agent integration"""
        # This would integrate with the google_drive_publisher agent
        # For now, return mock response
        return {
            'file_id': f"drive_id_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            'file_url': f"https://drive.google.com/file/d/mock_id/view",
            'folder': folder_name
        }

def main():
    """Demonstration of document conversion system"""
    print("[DOCUMENT CONVERSION] System Starting...")
    
    try:
        # Initialize converter
        converter = DocumentConverter()
        
        # Test with sample files
        test_files = [
            "system/orchestration/UNIVERSAL_ORCHESTRATOR_CHECKLIST.md",
            "system/sops/SOP_Token_Optimization_2025.md",
            "system/sops/SOP_Automated_Content_Refinement_2025.md",
            "system/core_tools/analysis_tools/enhanced_seo_extraction_report.md"
        ]
        
        converted_files = []
        
        for test_file in test_files:
            if os.path.exists(test_file):
                try:
                    output_path = converter.convert_markdown_to_docx(test_file, "system/reports/docx_exports")
                    converted_files.append(output_path)
                    print(f"[SUCCESS] Converted: {test_file}")
                except Exception as e:
                    print(f"[ERROR] Failed to convert {test_file}: {e}")
        
        # Create output directory if it doesn't exist
        os.makedirs("system/reports/docx_exports", exist_ok=True)
        
        print(f"\n[COMPLETE] Document conversion system ready")
        print(f"[FILES] {len(converted_files)} files converted")
        print(f"[LOCATION] system/reports/docx_exports/")
        
        # Integration with Google Drive
        drive_integration = GoogleDriveIntegration(converter)
        print(f"[INTEGRATION] Google Drive integration ready")
        
        return True
    
    except Exception as e:
        print(f"[ERROR] System initialization failed: {e}")
        return False

if __name__ == "__main__":
    main()