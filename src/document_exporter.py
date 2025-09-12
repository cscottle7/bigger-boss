"""
Document Export System
Convert markdown outputs to professional formats (.docx, PDF, Google Docs)
"""

import markdown
import re
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from datetime import datetime
import logging
import base64
import io

# Document export dependencies
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available. Install with: pip install python-docx")

try:
    import pdfkit
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("pdfkit not available. Install with: pip install pdfkit")

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False
    print("WeasyPrint not available. Install with: pip install weasyprint")

# Google API dependencies (optional)
try:
    from google.auth.transport.requests import Request
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from googleapiclient.discovery import build
    GOOGLE_API_AVAILABLE = True
except ImportError:
    GOOGLE_API_AVAILABLE = False
    print("Google API client not available. Install with: pip install google-api-python-client google-auth-oauthlib")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class MarkdownProcessor:
    """Process markdown content for document conversion"""
    
    def __init__(self):
        self.markdown_parser = markdown.Markdown(
            extensions=['tables', 'toc', 'codehilite', 'fenced_code', 'attr_list']
        )
    
    def parse_markdown_structure(self, md_content: str) -> Dict[str, Any]:
        """Parse markdown content into structured format"""
        lines = md_content.split('\n')
        structure = {
            'title': '',
            'sections': [],
            'tables': [],
            'code_blocks': [],
            'lists': [],
            'images': []
        }
        
        current_section = None
        in_code_block = False
        code_block_content = []
        
        for line in lines:
            line = line.strip()
            
            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End of code block
                    structure['code_blocks'].append({
                        'language': code_block_content[0] if code_block_content else '',
                        'content': '\n'.join(code_block_content[1:]) if len(code_block_content) > 1 else ''
                    })
                    code_block_content = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                    language = line[3:].strip()
                    code_block_content = [language]
                continue
            
            if in_code_block:
                code_block_content.append(line)
                continue
            
            # Parse headings
            if line.startswith('#'):
                level = len(line.split()[0])
                title = line[level:].strip()
                
                if level == 1:
                    structure['title'] = title
                else:
                    section = {
                        'level': level,
                        'title': title,
                        'content': [],
                        'subsections': []
                    }
                    
                    if level == 2:
                        structure['sections'].append(section)
                        current_section = section
                    elif level > 2 and current_section:
                        current_section['subsections'].append(section)
            
            # Parse tables
            elif '|' in line and line.count('|') >= 2:
                if not any(table for table in structure['tables'] if line in str(table)):
                    # Start of a new table
                    table_data = self._extract_table_from_position(lines, lines.index(line))
                    if table_data:
                        structure['tables'].append(table_data)
            
            # Parse lists
            elif line.startswith(('-', '*', '+')):
                list_item = line[1:].strip()
                if current_section:
                    if 'list_items' not in current_section:
                        current_section['list_items'] = []
                    current_section['list_items'].append(list_item)
                else:
                    structure['lists'].append(list_item)
            
            # Parse images
            elif '![' in line:
                image_match = re.search(r'!\[([^\]]*)\]\(([^\)]+)\)', line)
                if image_match:
                    structure['images'].append({
                        'alt': image_match.group(1),
                        'url': image_match.group(2)
                    })
            
            # Regular content
            elif line and current_section:
                current_section['content'].append(line)
        
        return structure
    
    def _extract_table_from_position(self, lines: List[str], start_pos: int) -> Optional[Dict]:
        """Extract table data starting from a specific position"""
        table_lines = []
        
        # Look for table lines starting from the current position
        for i in range(start_pos, len(lines)):
            line = lines[i].strip()
            if '|' in line and line.count('|') >= 2:
                table_lines.append(line)
            elif table_lines:  # End of table
                break
        
        if len(table_lines) < 2:
            return None
        
        # Parse table
        headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        
        # Skip separator line
        rows = []
        for line in table_lines[2:]:  # Skip header and separator
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells:
                    rows.append(cells)
        
        return {
            'headers': headers,
            'rows': rows
        }


class WordDocumentExporter:
    """Export content to Microsoft Word format"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word document export")
        
        self.doc = None
        self.processor = MarkdownProcessor()
    
    def export_to_docx(self, 
                      md_content: str, 
                      output_path: str, 
                      title: str = None,
                      author: str = "Autonomous Agentic Marketing System",
                      company: str = "Discover Web Solutions") -> str:
        """Export markdown content to Word document"""
        
        logger.info(f"Exporting to Word document: {output_path}")
        
        # Parse markdown structure
        structure = self.processor.parse_markdown_structure(md_content)
        
        # Create document
        self.doc = Document()
        
        # Set document properties
        properties = self.doc.core_properties
        properties.author = author
        properties.comments = f"Generated by {author} on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        if company:
            properties.category = company
        
        # Add title
        doc_title = title or structure.get('title', 'Marketing Analysis Report')
        title_para = self.doc.add_heading(doc_title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        self._add_document_metadata(author, company)
        
        # Add content
        self._add_sections(structure['sections'])
        
        # Add tables
        if structure['tables']:
            self.doc.add_heading('Data Tables', level=2)
            for table_data in structure['tables']:
                self._add_table(table_data)
        
        # Add code blocks
        if structure['code_blocks']:
            self.doc.add_heading('Technical Details', level=2)
            for code_block in structure['code_blocks']:
                self._add_code_block(code_block)
        
        # Save document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        
        logger.info(f"Word document saved: {output_path}")
        return str(output_path)
    
    def _add_document_metadata(self, author: str, company: str):
        """Add document metadata section"""
        meta_para = self.doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = meta_para.add_run(f"Generated by: {author}")
        run.italic = True
        meta_para.add_run('\n')
        
        if company:
            run = meta_para.add_run(f"Organization: {company}")
            run.italic = True
            meta_para.add_run('\n')
        
        run = meta_para.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        run.italic = True
        
        self.doc.add_paragraph()  # Add spacing
    
    def _add_sections(self, sections: List[Dict]):
        """Add sections to document"""
        for section in sections:
            # Add section heading
            self.doc.add_heading(section['title'], level=section['level'])
            
            # Add section content
            if section['content']:
                content_text = '\n'.join(section['content'])
                para = self.doc.add_paragraph()
                self._format_paragraph_content(para, content_text)
            
            # Add list items
            if 'list_items' in section:
                for item in section['list_items']:
                    para = self.doc.add_paragraph(item, style='List Bullet')
            
            # Add subsections
            if section['subsections']:
                self._add_sections(section['subsections'])
    
    def _format_paragraph_content(self, para, content: str):
        """Format paragraph content with markdown-style formatting"""
        # Handle bold text
        parts = re.split(r'(\*\*[^*]+\*\*)', content)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                # Regular text
                para.add_run(part)
    
    def _add_table(self, table_data: Dict):
        """Add table to document"""
        if not table_data['headers'] or not table_data['rows']:
            return
        
        # Create table
        table = self.doc.add_table(
            rows=len(table_data['rows']) + 1,
            cols=len(table_data['headers'])
        )
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(table_data['headers']):
            if i < len(header_cells):
                header_cells[i].text = header
                # Make header bold
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        
        # Add rows
        for row_idx, row_data in enumerate(table_data['rows']):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(row_cells):
                    row_cells[col_idx].text = str(cell_data)
        
        self.doc.add_paragraph()  # Add spacing after table
    
    def _add_code_block(self, code_block: Dict):
        """Add code block to document"""
        if code_block['language']:
            self.doc.add_heading(f"Code ({code_block['language']})", level=3)
        else:
            self.doc.add_heading("Code Block", level=3)
        
        para = self.doc.add_paragraph()
        run = para.add_run(code_block['content'])
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        
        # Add gray background (if possible)
        try:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F5F5F5')
            para._element.get_or_add_pPr().append(shading)
        except:
            pass


class PDFExporter:
    """Export content to PDF format"""
    
    def __init__(self):
        self.processor = MarkdownProcessor()
    
    def export_to_pdf(self, 
                     md_content: str, 
                     output_path: str,
                     title: str = None,
                     use_weasyprint: bool = True) -> str:
        """Export markdown content to PDF"""
        
        logger.info(f"Exporting to PDF: {output_path}")
        
        if use_weasyprint and WEASYPRINT_AVAILABLE:
            return self._export_with_weasyprint(md_content, output_path, title)
        elif PDF_AVAILABLE:
            return self._export_with_pdfkit(md_content, output_path, title)
        else:
            raise ImportError("No PDF export library available. Install weasyprint or pdfkit.")
    
    def _export_with_weasyprint(self, md_content: str, output_path: str, title: str) -> str:
        """Export using WeasyPrint"""
        # Convert markdown to HTML
        html_content = markdown.markdown(
            md_content,
            extensions=['tables', 'toc', 'codehilite', 'fenced_code']
        )
        
        # Create full HTML document
        full_html = self._create_html_document(html_content, title)
        
        # Create CSS for styling
        css_styles = self._create_pdf_styles()
        
        # Generate PDF
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        HTML(string=full_html).write_pdf(
            str(output_path),
            stylesheets=[CSS(string=css_styles)]
        )
        
        logger.info(f"PDF saved using WeasyPrint: {output_path}")
        return str(output_path)
    
    def _export_with_pdfkit(self, md_content: str, output_path: str, title: str) -> str:
        """Export using pdfkit"""
        # Convert markdown to HTML
        html_content = markdown.markdown(
            md_content,
            extensions=['tables', 'toc', 'codehilite', 'fenced_code']
        )
        
        # Create full HTML document
        full_html = self._create_html_document(html_content, title)
        
        # Generate PDF
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        options = {
            'page-size': 'A4',
            'margin-top': '0.75in',
            'margin-right': '0.75in',
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8",
            'no-outline': None
        }
        
        pdfkit.from_string(full_html, str(output_path), options=options)
        
        logger.info(f"PDF saved using pdfkit: {output_path}")
        return str(output_path)
    
    def _create_html_document(self, content: str, title: str) -> str:
        """Create full HTML document with styling"""
        doc_title = title or "Marketing Analysis Report"
        
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{doc_title}</title>
            <style>
                {self._create_pdf_styles()}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>{doc_title}</h1>
                <p class="generated-info">Generated by Autonomous Agentic Marketing System<br>
                Date: {datetime.now().strftime('%B %d, %Y')}</p>
            </div>
            <div class="content">
                {content}
            </div>
        </body>
        </html>
        """
    
    def _create_pdf_styles(self) -> str:
        """Create CSS styles for PDF"""
        return """
        body {
            font-family: 'Segoe UI', Arial, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            color: #333;
        }
        
        .header {
            text-align: center;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 2px solid #007acc;
        }
        
        .header h1 {
            color: #007acc;
            margin-bottom: 10px;
        }
        
        .generated-info {
            font-style: italic;
            color: #666;
        }
        
        .content {
            max-width: 800px;
            margin: 0 auto;
        }
        
        h1, h2, h3, h4, h5, h6 {
            color: #007acc;
            margin-top: 24px;
            margin-bottom: 12px;
        }
        
        h1 { font-size: 28px; }
        h2 { font-size: 24px; }
        h3 { font-size: 20px; }
        
        p { margin-bottom: 12px; }
        
        ul, ol {
            padding-left: 20px;
            margin-bottom: 12px;
        }
        
        li { margin-bottom: 6px; }
        
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }
        
        th, td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }
        
        th {
            background-color: #f2f2f2;
            font-weight: bold;
        }
        
        code {
            background-color: #f5f5f5;
            padding: 2px 4px;
            border-radius: 3px;
            font-family: Consolas, Monaco, monospace;
        }
        
        pre {
            background-color: #f5f5f5;
            padding: 12px;
            border-radius: 5px;
            overflow-x: auto;
        }
        
        blockquote {
            border-left: 4px solid #007acc;
            margin: 20px 0;
            padding-left: 16px;
            color: #666;
        }
        """


class GoogleDocsExporter:
    """Export content to Google Docs"""
    
    def __init__(self, credentials_file: str = None):
        if not GOOGLE_API_AVAILABLE:
            raise ImportError("Google API client libraries are required for Google Docs export")
        
        self.credentials_file = credentials_file
        self.service = None
        self.processor = MarkdownProcessor()
        
        # Scopes for Google Docs API
        self.SCOPES = ['https://www.googleapis.com/auth/documents']
    
    def authenticate(self, credentials_file: str = None) -> bool:
        """Authenticate with Google API"""
        try:
            creds = None
            token_file = 'token.json'
            
            # Load existing credentials
            if Path(token_file).exists():
                creds = Credentials.from_authorized_user_file(token_file, self.SCOPES)
            
            # If no valid credentials, request authorization
            if not creds or not creds.valid:
                if creds and creds.expired and creds.refresh_token:
                    creds.refresh(Request())
                else:
                    if not credentials_file:
                        logger.error("Google credentials file required for first-time authentication")
                        return False
                    
                    flow = InstalledAppFlow.from_client_secrets_file(
                        credentials_file, self.SCOPES)
                    creds = flow.run_local_server(port=0)
                
                # Save credentials for future use
                with open(token_file, 'w') as token:
                    token.write(creds.to_json())
            
            # Build service
            self.service = build('docs', 'v1', credentials=creds)
            logger.info("Google Docs authentication successful")
            return True
            
        except Exception as e:
            logger.error(f"Google Docs authentication failed: {str(e)}")
            return False
    
    def export_to_google_docs(self, 
                             md_content: str,
                             title: str = None,
                             folder_id: str = None) -> str:
        """Export markdown content to Google Docs"""
        
        if not self.service and not self.authenticate():
            raise Exception("Google Docs authentication required")
        
        logger.info(f"Creating Google Doc: {title}")
        
        # Create document
        doc_title = title or f"Marketing Analysis - {datetime.now().strftime('%Y-%m-%d')}"
        
        document = {
            'title': doc_title
        }
        
        doc = self.service.documents().create(body=document).execute()
        document_id = doc.get('documentId')
        
        # Parse markdown and add content
        structure = self.processor.parse_markdown_structure(md_content)
        
        requests = []
        
        # Add title and metadata
        requests.extend(self._create_title_requests(structure.get('title', doc_title)))
        
        # Add sections
        for section in structure['sections']:
            requests.extend(self._create_section_requests(section))
        
        # Add tables
        for table_data in structure['tables']:
            requests.extend(self._create_table_requests(table_data))
        
        # Execute all requests
        if requests:
            self.service.documents().batchUpdate(
                documentId=document_id,
                body={'requests': requests}
            ).execute()
        
        # Get document URL
        doc_url = f"https://docs.google.com/document/d/{document_id}/edit"
        
        logger.info(f"Google Doc created: {doc_url}")
        return doc_url
    
    def _create_title_requests(self, title: str) -> List[Dict]:
        """Create requests for document title"""
        return [
            {
                'insertText': {
                    'location': {'index': 1},
                    'text': f"{title}\n\n"
                }
            },
            {
                'updateParagraphStyle': {
                    'range': {
                        'startIndex': 1,
                        'endIndex': len(title) + 1
                    },
                    'paragraphStyle': {
                        'namedStyleType': 'TITLE',
                        'alignment': 'CENTER'
                    },
                    'fields': 'namedStyleType,alignment'
                }
            }
        ]
    
    def _create_section_requests(self, section: Dict) -> List[Dict]:
        """Create requests for a section"""
        requests = []
        
        # Add section heading
        heading_text = f"{section['title']}\n"
        requests.append({
            'insertText': {
                'location': {'index': 1},
                'text': heading_text
            }
        })
        
        # Style heading based on level
        style_map = {
            2: 'HEADING_1',
            3: 'HEADING_2',
            4: 'HEADING_3'
        }
        style = style_map.get(section['level'], 'HEADING_1')
        
        requests.append({
            'updateParagraphStyle': {
                'range': {
                    'startIndex': 1,
                    'endIndex': len(heading_text)
                },
                'paragraphStyle': {
                    'namedStyleType': style
                },
                'fields': 'namedStyleType'
            }
        })
        
        # Add section content
        if section['content']:
            content_text = '\n'.join(section['content']) + '\n\n'
            requests.append({
                'insertText': {
                    'location': {'index': 1},
                    'text': content_text
                }
            })
        
        # Add list items
        if 'list_items' in section:
            for item in section['list_items']:
                list_text = f"• {item}\n"
                requests.append({
                    'insertText': {
                        'location': {'index': 1},
                        'text': list_text
                    }
                })
        
        return requests
    
    def _create_table_requests(self, table_data: Dict) -> List[Dict]:
        """Create requests for a table"""
        requests = []
        
        if not table_data['headers'] or not table_data['rows']:
            return requests
        
        # Calculate table dimensions
        rows = len(table_data['rows']) + 1  # +1 for header
        cols = len(table_data['headers'])
        
        # Insert table
        requests.append({
            'insertTable': {
                'location': {'index': 1},
                'rows': rows,
                'columns': cols
            }
        })
        
        return requests


class DocumentExportManager:
    """Main manager for all document export operations"""
    
    def __init__(self):
        self.word_exporter = None
        self.pdf_exporter = None
        self.gdocs_exporter = None
    
    def export_markdown(self,
                       md_content: str,
                       output_formats: List[str],
                       base_filename: str = None,
                       output_dir: str = "exports",
                       title: str = None,
                       **kwargs) -> Dict[str, str]:
        """
        Export markdown to multiple formats
        
        Args:
            md_content: Markdown content to export
            output_formats: List of formats ('docx', 'pdf', 'google_docs')
            base_filename: Base filename (without extension)
            output_dir: Output directory
            title: Document title
            **kwargs: Additional export options
        
        Returns:
            Dictionary mapping format to output path/URL
        """
        
        results = {}
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if not base_filename:
            base_filename = f"marketing_analysis_{timestamp}"
        
        # Ensure output directory exists
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        for format_type in output_formats:
            try:
                if format_type.lower() == 'docx':
                    if not self.word_exporter:
                        self.word_exporter = WordDocumentExporter()
                    
                    output_path = Path(output_dir) / f"{base_filename}.docx"
                    result = self.word_exporter.export_to_docx(
                        md_content, 
                        str(output_path),
                        title=title,
                        **kwargs
                    )
                    results['docx'] = result
                
                elif format_type.lower() == 'pdf':
                    if not self.pdf_exporter:
                        self.pdf_exporter = PDFExporter()
                    
                    output_path = Path(output_dir) / f"{base_filename}.pdf"
                    result = self.pdf_exporter.export_to_pdf(
                        md_content,
                        str(output_path),
                        title=title,
                        **kwargs
                    )
                    results['pdf'] = result
                
                elif format_type.lower() == 'google_docs':
                    if not self.gdocs_exporter:
                        credentials_file = kwargs.get('google_credentials')
                        self.gdocs_exporter = GoogleDocsExporter(credentials_file)
                    
                    result = self.gdocs_exporter.export_to_google_docs(
                        md_content,
                        title=title,
                        **kwargs
                    )
                    results['google_docs'] = result
                
                else:
                    logger.warning(f"Unsupported export format: {format_type}")
                    
            except Exception as e:
                logger.error(f"Error exporting to {format_type}: {str(e)}")
                results[format_type] = f"Error: {str(e)}"
        
        return results
    
    def get_available_formats(self) -> List[str]:
        """Get list of available export formats"""
        formats = []
        
        if DOCX_AVAILABLE:
            formats.append('docx')
        
        if PDF_AVAILABLE or WEASYPRINT_AVAILABLE:
            formats.append('pdf')
        
        if GOOGLE_API_AVAILABLE:
            formats.append('google_docs')
        
        return formats


# Example usage
if __name__ == "__main__":
    # Sample markdown content
    sample_md = """
# Marketing Analysis Report

## Executive Summary

This report provides comprehensive analysis of the website performance and competitive landscape.

### Key Findings

- Website performance needs improvement
- Content strategy should focus on **mobile optimization**
- Competitive analysis shows opportunities in social media

## Technical Analysis

| Metric | Current | Target |
|--------|---------|--------|
| Load Time | 3.2s | <2.0s |
| Mobile Score | 65 | >90 |

### Recommendations

1. Optimize image loading
2. Implement caching
3. Improve mobile responsiveness

## Code Example

```python
def optimize_images():
    return "compressed_images"
```

## Conclusion

Implementation of these recommendations will improve overall site performance by 40%.
    """
    
    # Test exports
    exporter = DocumentExportManager()
    available_formats = exporter.get_available_formats()
    
    print(f"Available formats: {available_formats}")
    
    if available_formats:
        results = exporter.export_markdown(
            sample_md,
            available_formats,
            base_filename="test_report",
            title="Test Marketing Analysis"
        )
        
        print("Export results:")
        for format_type, result in results.items():
            print(f"  {format_type}: {result}")
    else:
        print("No export libraries available. Install python-docx, weasyprint, or google-api-python-client")